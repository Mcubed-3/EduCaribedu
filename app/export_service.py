from __future__ import annotations

import html
import re
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Pt
from playwright.sync_api import sync_playwright

EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)

BOLD_LABELS = ("Formula:", "Substitute:", "Simplify:", "Answer:", "Method:", "Worked Example:")


def _clean_text(value: str) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    return text.strip()


def _is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    return stripped.endswith(":") and len(stripped) <= 80 and not _is_numbered_item(stripped)


def _is_list_item(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("- ") or stripped.startswith("• ")


def _is_numbered_item(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    return re.match(r"^\d+\.\s", stripped) is not None


def _starts_with_bold_label(line: str) -> bool:
    stripped = line.strip()
    return any(stripped.startswith(label) for label in BOLD_LABELS)


def _split_label_line(line: str):
    stripped = line.strip()
    for label in BOLD_LABELS:
        if stripped.startswith(label):
            return label, stripped[len(label):].strip()
    return None, stripped


def _is_table_title(line: str) -> bool:
    return line.strip().startswith("Table:")


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return "|" in stripped and len([c for c in stripped.split("|")]) >= 2


def _parse_table_block(lines, start_index):
    table_title = lines[start_index].strip().replace("Table:", "").strip()
    rows = []
    i = start_index + 1

    while i < len(lines):
        line = lines[i].strip()
        if _is_table_row(line):
            rows.append([cell.strip() for cell in line.split("|")])
            i += 1
        else:
            break

    max_cols = max((len(r) for r in rows), default=0)
    padded_rows = [r + [""] * (max_cols - len(r)) for r in rows]

    return table_title, padded_rows, i


def _build_html_from_text(title: str, content: str) -> str:
    safe_title = html.escape(title or "Export")
    lines = _clean_text(content).splitlines()

    html_lines = [
        "<html>",
        "<head>",
        '<meta charset="utf-8" />',
        f"<title>{safe_title}</title>",
        "<style>",
        "body { font-family: Arial, sans-serif; color: #14324a; background: white; margin: 0; padding: 28px; line-height: 1.55; }",
        ".wrapper { max-width: 900px; margin: 0 auto; }",
        "h1 { font-size: 22px; margin: 0 0 18px 0; }",
        "h2 { font-size: 18px; margin: 20px 0 10px 0; }",
        "p { margin: 0 0 10px 0; white-space: pre-wrap; }",
        ".list-item { margin: 0 0 8px 0; white-space: pre-wrap; }",
        ".number-item { margin: 0 0 10px 0; white-space: pre-wrap; }",
        ".label-line { margin: 0 0 10px 0; white-space: pre-wrap; }",
        ".spacer { height: 10px; }",
        "table { border-collapse: collapse; width: 100%; margin: 0 0 14px 0; }",
        "th, td { border: 1px solid #d7e4f1; padding: 8px 10px; text-align: left; font-size: 13px; }",
        "th { background: #eef5ff; font-weight: 700; }",
        "</style>",
        "</head>",
        "<body>",
        '<div class="wrapper">',
    ]

    if title:
        html_lines.append(f"<h1>{safe_title}</h1>")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            html_lines.append('<div class="spacer"></div>')
            i += 1
            continue

        if _is_table_title(stripped):
            table_title, rows, new_i = _parse_table_block(lines, i)
            if table_title:
                html_lines.append(f"<h2>{html.escape(table_title)}</h2>")
            if rows:
                html_lines.append("<table>")
                for row_idx, row in enumerate(rows):
                    tag = "th" if row_idx == 0 else "td"
                    html_lines.append(
                        "<tr>" + "".join(f"<{tag}>{html.escape(cell)}</{tag}>" for cell in row) + "</tr>"
                    )
                html_lines.append("</table>")
            i = new_i
            continue

        if _starts_with_bold_label(stripped):
            label, rest = _split_label_line(stripped)
            html_lines.append(
                f'<p class="label-line"><strong>{html.escape(label)}</strong> {html.escape(rest)}</p>'
            )
            i += 1
            continue

        escaped = html.escape(line)

        if _is_heading(stripped):
            html_lines.append(f"<h2>{html.escape(stripped[:-1])}</h2>")
        elif _is_numbered_item(stripped):
            html_lines.append(f'<p class="number-item">{escaped}</p>')
        elif _is_list_item(stripped):
            html_lines.append(f'<p class="list-item">{escaped}</p>')
        else:
            html_lines.append(f"<p>{escaped}</p>")

        i += 1

    html_lines.extend([
        "</div>",
        "</body>",
        "</html>",
    ])

    return "\n".join(html_lines)


def export_to_docx(title: str, content: str):
    file_name = f"{uuid.uuid4()}.docx"
    file_path = EXPORT_DIR / file_name

    doc = Document()
    doc.add_heading(title or "Export", level=0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    lines = _clean_text(content).splitlines()
    i = 0

    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if _is_table_title(stripped):
            table_title, rows, new_i = _parse_table_block(lines, i)

            if table_title:
                p = doc.add_paragraph()
                run = p.add_run(table_title)
                run.bold = True

            if rows:
                cols = len(rows[0])
                table = doc.add_table(rows=1, cols=cols)
                table.style = "Table Grid"

                for j, cell in enumerate(rows[0]):
                    run = table.rows[0].cells[j].paragraphs[0].add_run(cell)
                    run.bold = True

                for row in rows[1:]:
                    cells = table.add_row().cells
                    for j, cell in enumerate(row):
                        cells[j].text = cell

            i = new_i
            continue

        if _is_heading(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            i += 1
            continue

        if _starts_with_bold_label(stripped):
            label, rest = _split_label_line(stripped)
            p = doc.add_paragraph()
            label_run = p.add_run(label)
            label_run.bold = True
            if rest:
                p.add_run(f" {rest}")
            i += 1
            continue

        if _is_list_item(stripped):
            doc.add_paragraph(stripped, style="List Bullet")
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    doc.save(file_path)
    return file_path


def export_to_pdf(title: str, content: str):
    file_name = f"{uuid.uuid4()}.pdf"
    file_path = EXPORT_DIR / file_name

    full_html = _build_html_from_text(title=title or "Export", content=content or "")

    with sync_playwright() as p:
        browser = p.chromium.launch(args=["--no-sandbox"])
        page = browser.new_page()
        page.set_content(full_html, wait_until="load")
        page.pdf(
            path=str(file_path),
            format="A4",
            print_background=True,
            margin={
                "top": "15mm",
                "bottom": "15mm",
                "left": "15mm",
                "right": "15mm",
            },
        )
        browser.close()

    return file_path